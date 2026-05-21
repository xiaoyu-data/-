import os
import shutil
import json
from pathlib import Path
from typing import Any


class BaseTool:
    def __init__(self):
        self.name = self.__class__.__name__
        self.description = ""

    def run(self, **kwargs) -> str:
        raise NotImplementedError


class FileTools(BaseTool):
    def __init__(self):
        super().__init__()
        self.description = "文件系统操作：列出文件、读写文本、创建目录、复制移动文件、删除文件"

    def list_files(self, path: str = ".", pattern: str = "*") -> str:
        p = Path(path)
        if not p.exists():
            return f"路径不存在: {path}"
        files = list(p.glob(pattern))
        if not files:
            return f"目录 '{path}' 下没有匹配 '{pattern}' 的文件"
        result = []
        for f in sorted(files):
            ftype = "[DIR]" if f.is_dir() else "[FILE]"
            size = f.stat().st_size if f.is_file() else 0
            result.append(f"  {ftype} {f.name}  ({self._format_size(size)})")
        return "\n".join(result)

    def read_text(self, filepath: str) -> str:
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()
        except FileNotFoundError:
            return f"文件不存在: {filepath}"
        except UnicodeDecodeError:
            try:
                with open(filepath, "r", encoding="gbk") as f:
                    return f.read()
            except Exception as e:
                return f"读取失败: {e}"

    def write_text(self, filepath: str, content: str) -> str:
        try:
            Path(filepath).parent.mkdir(parents=True, exist_ok=True)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(content)
            return f"已写入: {filepath} ({len(content)} 字符)"
        except Exception as e:
            return f"写入失败: {e}"

    def create_directory(self, path: str) -> str:
        try:
            Path(path).mkdir(parents=True, exist_ok=True)
            return f"已创建目录: {path}"
        except Exception as e:
            return f"创建目录失败: {e}"

    def copy_file(self, src: str, dst: str) -> str:
        try:
            Path(dst).parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(src, dst)
            return f"已复制: {src} -> {dst}"
        except Exception as e:
            return f"复制失败: {e}"

    def delete_file(self, filepath: str) -> str:
        try:
            os.remove(filepath)
            return f"已删除: {filepath}"
        except FileNotFoundError:
            return f"文件不存在: {filepath}"
        except Exception as e:
            return f"删除失败: {e}"

    @staticmethod
    def _format_size(size: int) -> str:
        for unit in ["B", "KB", "MB", "GB"]:
            if size < 1024:
                return f"{size:.1f}{unit}"
            size /= 1024
        return f"{size:.1f}TB"

    def run(self, action: str = "", **kwargs) -> str:
        actions = {
            "list": self.list_files,
            "read": self.read_text,
            "write": self.write_text,
            "mkdir": self.create_directory,
            "copy": self.copy_file,
            "delete": self.delete_file,
        }
        fn = actions.get(action)
        if fn:
            return fn(**kwargs)
        return f"未知的文件操作: {action}，支持: {list(actions.keys())}"


class WordTools(BaseTool):
    def __init__(self):
        super().__init__()
        self.description = "Word 文档处理：创建文档、读取内容、提取文本、替换文字、添加表格"

    def create_doc(self, filepath: str, content: str = "", title: str = "新文档") -> str:
        try:
            from docx import Document
        except ImportError:
            return "请安装 python-docx: pip install python-docx"

        doc = Document()
        doc.add_heading(title, level=1)
        for para in content.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.save(filepath)
        return f"已创建 Word 文档: {filepath}"

    def read_doc(self, filepath: str) -> str:
        try:
            from docx import Document
        except ImportError:
            return "请安装 python-docx: pip install python-docx"

        try:
            doc = Document(filepath)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            return "\n".join(parts)
        except FileNotFoundError:
            return f"文件不存在: {filepath}"
        except Exception as e:
            return f"读取 Word 文档失败: {e}"

    def extract_tables(self, filepath: str) -> str:
        try:
            from docx import Document
        except ImportError:
            return "请安装 python-docx: pip install python-docx"

        try:
            doc = Document(filepath)
            results = []
            for i, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    rows.append(" | ".join(cells))
                results.append(f"--- 表格 {i + 1} ---\n" + "\n".join(rows))
            return "\n\n".join(results) if results else "文档中没有表格"
        except Exception as e:
            return f"提取表格失败: {e}"

    def replace_text(self, filepath: str, old_text: str, new_text: str, output: str = "") -> str:
        try:
            from docx import Document
        except ImportError:
            return "请安装 python-docx: pip install python-docx"

        try:
            doc = Document(filepath)
            count = 0
            for para in doc.paragraphs:
                if old_text in para.text:
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            count += 1
            output = output or filepath
            doc.save(output)
            return f"已在 {count} 处替换 '{old_text}' -> '{new_text}'，保存至: {output}"
        except Exception as e:
            return f"替换失败: {e}"

    def run(self, action: str = "", **kwargs) -> str:
        actions = {
            "create": self.create_doc,
            "read": self.read_doc,
            "tables": self.extract_tables,
            "replace": self.replace_text,
        }
        fn = actions.get(action)
        if fn:
            return fn(**kwargs)
        return f"未知的 Word 操作: {action}，支持: {list(actions.keys())}"


class ExcelTools(BaseTool):
    def __init__(self):
        super().__init__()
        self.description = "Excel 表格处理：创建工作簿、读取数据、写入数据、提取表信息、数据分析"

    def create_workbook(self, filepath: str, sheet_data: dict = None) -> str:
        try:
            from openpyxl import Workbook
        except ImportError:
            return "请安装 openpyxl: pip install openpyxl"

        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Sheet1"
            if sheet_data:
                for row_idx, row in enumerate(sheet_data.get("data", []), 1):
                    for col_idx, value in enumerate(row, 1):
                        ws.cell(row=row_idx, column=col_idx, value=value)
            wb.save(filepath)
            return f"已创建 Excel 工作簿: {filepath}"
        except Exception as e:
            return f"创建失败: {e}"

    def read_sheet(self, filepath: str, sheet_name: str = "") -> str:
        try:
            from openpyxl import load_workbook
        except ImportError:
            return "请安装 openpyxl: pip install openpyxl"

        try:
            wb = load_workbook(filepath, data_only=True)
            ws = wb[sheet_name] if sheet_name else wb.active
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append(" | ".join(str(c) if c is not None else "" for c in row))
            return "\n".join(rows) if rows else "工作表为空"
        except Exception as e:
            return f"读取 Excel 失败: {e}"

    def write_data(self, filepath: str, data: list, sheet_name: str = "Sheet1") -> str:
        try:
            from openpyxl import load_workbook, Workbook
        except ImportError:
            return "请安装 openpyxl: pip install openpyxl"

        try:
            if Path(filepath).exists():
                wb = load_workbook(filepath)
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                else:
                    ws = wb.create_sheet(sheet_name)
            else:
                wb = Workbook()
                ws = wb.active
                ws.title = sheet_name

            for row_idx, row in enumerate(data, 1):
                for col_idx, value in enumerate(row, 1):
                    ws.cell(row=row_idx, column=col_idx, value=value)
            wb.save(filepath)
            return f"已向 '{sheet_name}' 写入 {len(data)} 行数据"
        except Exception as e:
            return f"写入失败: {e}"

    def analyze(self, filepath: str) -> str:
        try:
            import pandas as pd
        except ImportError:
            return "请安装 pandas: pip install pandas"

        try:
            df = pd.read_excel(filepath)
            info = [
                f"行数: {len(df)}, 列数: {len(df.columns)}",
                f"列名: {list(df.columns)}",
                f"数值列统计:\n{df.describe().to_string()}",
            ]
            return "\n".join(info)
        except Exception as e:
            return f"分析失败: {e}"

    def run(self, action: str = "", **kwargs) -> str:
        actions = {
            "create": self.create_workbook,
            "read": self.read_sheet,
            "write": self.write_data,
            "analyze": self.analyze,
        }
        fn = actions.get(action)
        if fn:
            return fn(**kwargs)
        return f"未知的 Excel 操作: {action}，支持: {list(actions.keys())}"


class PDFTools(BaseTool):
    def __init__(self):
        super().__init__()
        self.description = "PDF 文档处理：提取文本、提取表格、获取文档信息"

    def extract_text(self, filepath: str) -> str:
        try:
            import pdfplumber
        except ImportError:
            return "请安装 pdfplumber: pip install pdfplumber"

        try:
            with pdfplumber.open(filepath) as pdf:
                parts = []
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        parts.append(f"--- 第 {i + 1} 页 ---\n{text}")
                return "\n\n".join(parts) if parts else "PDF 中没有可提取的文字"
        except Exception as e:
            return f"提取 PDF 文本失败: {e}"

    def extract_tables(self, filepath: str) -> str:
        try:
            import pdfplumber
        except ImportError:
            return "请安装 pdfplumber: pip install pdfplumber"

        try:
            with pdfplumber.open(filepath) as pdf:
                results = []
                for i, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    for j, table in enumerate(tables):
                        rows = []
                        for row in table:
                            rows.append(" | ".join(str(c) if c else "" for c in row))
                        results.append(
                            f"--- 第 {i + 1} 页 表格 {j + 1} ---\n" + "\n".join(rows)
                        )
                return "\n\n".join(results) if results else "PDF 中没有表格"
        except Exception as e:
            return f"提取 PDF 表格失败: {e}"

    def get_info(self, filepath: str) -> str:
        try:
            import pdfplumber
        except ImportError:
            return "请安装 pdfplumber: pip install pdfplumber"

        try:
            with pdfplumber.open(filepath) as pdf:
                info = [
                    f"总页数: {len(pdf.pages)}",
                ]
                if pdf.metadata:
                    for k, v in pdf.metadata.items():
                        if v:
                            info.append(f"{k}: {v}")
                return "\n".join(info)
        except Exception as e:
            return f"获取 PDF 信息失败: {e}"

    def run(self, action: str = "", **kwargs) -> str:
        actions = {
            "extract_text": self.extract_text,
            "extract_tables": self.extract_tables,
            "get_info": self.get_info,
        }
        fn = actions.get(action)
        if fn:
            return fn(**kwargs)
        return f"未知的 PDF 操作: {action}，支持: {list(actions.keys())}"
